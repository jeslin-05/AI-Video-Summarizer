from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from downloader import download_audio
from transcriber import transcribe_audio
from summarizer import summarize_text
from translator import translate_text
from fastapi.responses import FileResponse
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document
import mysql.connector
import os
import uuid

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- DATABASE CONNECTION (XAMPP) ---
def get_db_connection():
    return mysql.connector.connect(
        host="localhost",
        user="root",
        password="", 
        database="video_summarizer"
    )

# --- UNICODE FONT REGISTRATION ---
FONT_NAME = 'Helvetica'
try:
    font_path = "arial.ttf"
    if not os.path.exists(font_path):
        font_path = "C:/Windows/Fonts/arial.ttf"
    if os.path.exists(font_path):
        pdfmetrics.registerFont(TTFont('UnicodeArial', font_path))
        FONT_NAME = 'UnicodeArial'
except Exception as e:
    print(f"Font registration failed: {e}")

# --- MODELS ---
class LoginRequest(BaseModel):
    username: str
    password: str

class SignupRequest(BaseModel):
    username: str
    password: str
    mobile: str

class VideoRequest(BaseModel):
    url: str

class TranslateRequest(BaseModel):
    text: str
    target_lang: str

# --- AUTH ENDPOINTS ---

@app.post("/signup")
def signup(request: SignupRequest):
    db = get_db_connection()
    cursor = db.cursor()
    try:
        cursor.execute("INSERT INTO users (username, password, mobile_number) VALUES (%s, %s, %s)", 
                       (request.username, request.password, request.mobile))
        db.commit()
        return {"status": "success", "message": "User registered"}
    except mysql.connector.Error:
        raise HTTPException(status_code=400, detail="Username already exists")
    finally:
        db.close()

@app.post("/login")
def login(request: LoginRequest):
    db = get_db_connection()
    cursor = db.cursor()
    cursor.execute("SELECT * FROM users WHERE username = %s AND password = %s", 
                   (request.username, request.password))
    user = cursor.fetchone()
    db.close()
    if user:
        return {"status": "success", "username": request.username}
    raise HTTPException(status_code=401, detail="Invalid username or password")

# --- SUMMARIZE ENDPOINT (FIXED LOGIC) ---

@app.post("/summarize")
def summarize_video(request: VideoRequest):
    try:
        audio_file = download_audio(request.url)
        # transcriber returns (transcript_text, detected_lang_code)
        transcript, detected_lang = transcribe_audio(audio_file)
        
        # FIX: Only translate if the detected language is definitely NOT English
        if detected_lang and detected_lang.lower() != "en":
            from transformers import pipeline
            # Use 1024 max_length to fix the "512 limit" error in your logs
            translator = pipeline("translation", model="Helsinki-NLP/opus-mt-mul-en")
            translation_result = translator(transcript, max_length=1024, truncation=True)
            transcript = translation_result[0]["translation_text"]
        
        summary = summarize_text(transcript)
        
        if os.path.exists(audio_file):
            os.remove(audio_file)
            
        return {
            "status": "success", 
            "summary": summary, 
            "detected_language": detected_lang if detected_lang else "en"
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}

@app.post("/translate")
def translate_summary(request: TranslateRequest):
    try:
        translated = translate_text(request.text, request.target_lang)
        return {"translated_text": translated}
    except Exception as e:
        return {"status": "error", "message": str(e)}

@app.post("/export/pdf")
def export_pdf(request: TranslateRequest):
    filename = f"{uuid.uuid4()}.pdf"
    doc = SimpleDocTemplate(filename)
    styles = getSampleStyleSheet()
    unicode_style = ParagraphStyle('UnicodeStyle', parent=styles['Normal'], fontName=FONT_NAME, fontSize=12, leading=18, wordWrap='CJK')
    elements = [Paragraph("<b>AI Video Summary</b>", styles['Title']), Spacer(1, 20), Paragraph(request.text, unicode_style)]
    doc.build(elements)
    return FileResponse(filename, media_type="application/pdf", filename="summary.pdf")

@app.post("/export/doc")
def export_doc(request: TranslateRequest):
    filename = f"{uuid.uuid4()}.docx"
    document = Document()
    style = document.styles['Normal']
    style.font.name = 'Arial'
    document.add_heading('AI Video Summary', 0)
    document.add_paragraph(request.text)
    document.save(filename)
    return FileResponse(filename, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="summary.docx")